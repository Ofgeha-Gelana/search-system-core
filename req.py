# import re
# import string
# import numpy as np
# import pandas as pd
# import nltk
# import spacy
# import streamlit as st
# from sklearn.feature_extraction.text import TfidfVectorizer
# from nltk.corpus import stopwords
# from nltk.stem import WordNetLemmatizer
# import fitz  # PyMuPDF for PDF text extraction
# from docx import Document  # python-docx for Word documents
# from bs4 import BeautifulSoup
# import requests
# import io  # To handle file-like objects for PDF processing

# # --- NLTK Setup ---
# nltk.download("stopwords")
# nltk.download("punkt")
# nltk.download("wordnet")

# # --- Initialization ---
# lemmatizer = WordNetLemmatizer()

# # --- Text Extraction Functions ---
# def extract_pdf_text(file):
#     """Extract text from a PDF file."""
#     doc = fitz.open(stream=file.read(), filetype="pdf")
#     text = "".join([page.get_text() for page in doc])
#     return text.strip()

# def extract_word_text(file):
#     """Extract text from a Word document."""
#     doc = Document(file)
#     text = "\n".join([para.text for para in doc.paragraphs])
#     return text.strip()

# def extract_website_text(url):
#     """Extract text from a website URL."""
#     response = requests.get(url)
#     soup = BeautifulSoup(response.content, "html.parser")
#     paragraphs = soup.find_all("p")
#     text = " ".join([para.get_text() for para in paragraphs])
#     return text.strip()

# # --- Text Preprocessing ---
# def preprocess_text(docs):
#     """Clean and lemmatize text data."""
#     cleaned_docs = []
#     for doc in docs:
#         doc = re.sub(r'[^\x00-\x7F]+', ' ', doc)  # Remove non-ASCII characters
#         doc = re.sub(r'@\w+', '', doc)  # Remove mentions
#         doc = doc.lower()  # Convert to lowercase
#         doc = re.sub(r'[%s]' % re.escape(string.punctuation), ' ', doc)  # Remove punctuation
#         doc = re.sub(r'\d+', '', doc)  # Remove numbers
#         doc = re.sub(r'\s{2,}', ' ', doc)  # Remove extra spaces
#         lemmatized_doc = " ".join([lemmatizer.lemmatize(word) for word in doc.split()])
#         cleaned_docs.append(lemmatized_doc)
#     return cleaned_docs

# # --- Utility Functions ---
# def split_into_sentences(text):
#     """Split text into sentences."""
#     return nltk.sent_tokenize(text)

# def get_similar_sentences(query, vectorizer, df, top_n=5):
#     """Find similar sentences based on cosine similarity."""
#     query = preprocess_text([query])  # Preprocess the query
#     query_vec = vectorizer.transform(query).toarray().flatten()

#     similarity_scores = {}
#     for i, sentence in enumerate(new_docs):
#         sentence_vec = X[i].toarray().flatten()
#         numerator = np.dot(query_vec, sentence_vec)
#         denominator = np.linalg.norm(query_vec) * np.linalg.norm(sentence_vec)
#         similarity = numerator / denominator if denominator != 0 else 0
#         similarity_scores[i] = similarity

#     # Sort by similarity scores in descending order
#     sorted_scores = sorted(similarity_scores.items(), key=lambda x: x[1], reverse=True)[:top_n]

#     results = []
#     for idx, score in sorted_scores:
#         if score > 0:  # Only include sentences with positive similarity
#             results.append((f"Sentence {idx + 1}", sentences[idx], score))
#     return results

# # --- Streamlit UI ---
# # Load custom styling
# st.markdown(
#     """
#     <style>
#     .title {
#         font-size: 36px;
#         color: #4CAF50;
#         text-align: center;
#         font-weight: bold;
#     }
#     .subtitle {
#         font-size: 18px;
#         text-align: center;
#         color: #555;
#     }
#     .menu-btn {
#         background-color: #007BFF;
#         color: white;
#         border: none;
#         border-radius: 5px;
#         padding: 10px 20px;
#         margin: 10px;
#         font-size: 16px;
#         cursor: pointer;
#         text-align: center;
#     }
#     .menu-btn:hover {
#         background-color: #0056b3;
#     }
#     .results {
#         margin-top: 20px;
#     }
#     .result-item {
#         background-color: #f8f9fa;
#         padding: 10px;
#         border-radius: 8px;
#         margin-bottom: 10px;
#         box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);
#     }
#     </style>
#     """,
#     unsafe_allow_html=True,
# )

# st.markdown('<div class="title">Document Search System</div>', unsafe_allow_html=True)
# st.markdown('<div class="subtitle">Upload a document or provide a URL to perform a search query.</div>', unsafe_allow_html=True)

# # --- Menu Options ---
# st.markdown('<div style="text-align: center;">', unsafe_allow_html=True)
# option = st.radio("Choose Input Method:", ["Upload File", "Enter URL"], horizontal=True)
# st.markdown('</div>', unsafe_allow_html=True)

# # --- Input Handling ---
# document_text = ""
# if option == "Upload File":
#     uploaded_file = st.file_uploader("Upload a PDF or Word document", type=["pdf", "docx"])
#     if uploaded_file:
#         if uploaded_file.type == "application/pdf":
#             document_text = extract_pdf_text(uploaded_file)
#         elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#             document_text = extract_word_text(uploaded_file)
#         else:
#             st.warning("Unsupported file format. Please upload a PDF or Word document.")
# elif option == "Enter URL":
#     url_input = st.text_input("Enter website URL:")
#     if url_input:
#         document_text = extract_website_text(url_input)

# # --- Text Processing and Query ---
# if document_text:
#     sentences = split_into_sentences(document_text)
#     new_docs = preprocess_text(sentences)

#     vectorizer = TfidfVectorizer(
#         analyzer="word",
#         ngram_range=(1, 2),
#         min_df=0.01,
#         max_df=0.85,
#         stop_words="english",
#     )
#     X = vectorizer.fit_transform(new_docs)

#     query = st.text_input("Enter search query:")
#     top_n_results = st.slider("Number of top results to display:", min_value=1, max_value=20, value=5, step=1)

#     if query:
#         results = get_similar_sentences(query, vectorizer, df=None, top_n=top_n_results)
#         if results:
#             st.markdown('<div class="results">', unsafe_allow_html=True)
#             for title, sentence, score in results:
#                 st.markdown(
#                     f'<div class="result-item"><strong>{title}</strong><br><em>Similarity Score: {score:.4f}</em><br>{sentence}</div>',
#                     unsafe_allow_html=True,
#                 )
#             st.markdown('</div>', unsafe_allow_html=True)
#         else:
#             st.info("No similar sentences found.")
# else:
#     st.info("Upload a document or enter a website URL to begin.")





# import re
# import string
# import numpy as np
# import pandas as pd
# import nltk
# import spacy
# import streamlit as st
# from sklearn.feature_extraction.text import TfidfVectorizer
# from nltk.corpus import stopwords
# from nltk.stem import WordNetLemmatizer
# import fitz  # PyMuPDF for PDF text extraction
# from docx import Document  # python-docx for Word documents
# from bs4 import BeautifulSoup
# import requests
# import io  # To handle file-like objects for PDF processing

# # --- NLTK Setup ---
# nltk.download("stopwords")
# nltk.download("punkt")
# nltk.download("wordnet")

# # --- Initialization ---
# lemmatizer = WordNetLemmatizer()

# # --- Text Extraction Functions ---
# def extract_pdf_text(file):
#     """Extract text from a PDF file."""
#     doc = fitz.open(stream=file.read(), filetype="pdf")
#     text = "".join([page.get_text() for page in doc])
#     return text.strip()

# def extract_word_text(file):
#     """Extract text from a Word document."""
#     doc = Document(file)
#     text = "\n".join([para.text for para in doc.paragraphs])
#     return text.strip()

# def extract_website_text(url):
#     """Extract text from a website URL."""
#     response = requests.get(url)
#     soup = BeautifulSoup(response.content, "html.parser")
#     paragraphs = soup.find_all("p")
#     text = " ".join([para.get_text() for para in paragraphs])
#     return text.strip()

# # --- Text Preprocessing ---
# def preprocess_text(docs):
#     """Clean and lemmatize text data."""
#     cleaned_docs = []
#     for doc in docs:
#         doc = re.sub(r'[^\x00-\x7F]+', ' ', doc)  # Remove non-ASCII characters
#         doc = re.sub(r'@\w+', '', doc)  # Remove mentions
#         doc = doc.lower()  # Convert to lowercase
#         doc = re.sub(r'[%s]' % re.escape(string.punctuation), ' ', doc)  # Remove punctuation
#         doc = re.sub(r'\d+', '', doc)  # Remove numbers
#         doc = re.sub(r'\s{2,}', ' ', doc)  # Remove extra spaces
#         lemmatized_doc = " ".join([lemmatizer.lemmatize(word) for word in doc.split()])
#         cleaned_docs.append(lemmatized_doc)
#     return cleaned_docs

# # --- Utility Functions ---
# def split_into_sentences(text):
#     """Split text into sentences."""
#     return nltk.sent_tokenize(text)

# def get_similar_sentences(query, vectorizer, df, top_n=5):
#     """Find similar sentences based on cosine similarity."""
#     query = preprocess_text([query])  # Preprocess the query
#     query_vec = vectorizer.transform(query).toarray().flatten()

#     similarity_scores = {}
#     for i, sentence in enumerate(new_docs):
#         sentence_vec = X[i].toarray().flatten()
#         numerator = np.dot(query_vec, sentence_vec)
#         denominator = np.linalg.norm(query_vec) * np.linalg.norm(sentence_vec)
#         similarity = numerator / denominator if denominator != 0 else 0
#         similarity_scores[i] = similarity

#     # Sort by similarity scores in descending order
#     sorted_scores = sorted(similarity_scores.items(), key=lambda x: x[1], reverse=True)[:top_n]

#     results = []
#     for idx, score in sorted_scores:
#         if score > 0:  # Only include sentences with positive similarity
#             results.append((f"Sentence {idx + 1}", sentences[idx], score))
#     return results

# # --- Streamlit UI ---
# st.title("Document Search System")
# st.subheader("Upload a document or provide a URL to perform a search query.")

# # File upload for PDF or Word documents
# uploaded_file = st.file_uploader("Upload a PDF or Word document", type=["pdf", "docx"], key="file_uploader")

# # URL input for scraping text
# url_input = st.text_input("Enter website URL (for scraping)", key="url_input")

# # Initialize document text
# document_text = ""

# # Extract text from the uploaded file or URL
# if uploaded_file:
#     if uploaded_file.type == "application/pdf":
#         document_text = extract_pdf_text(uploaded_file)
#     elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#         document_text = extract_word_text(uploaded_file)
#     else:
#         st.warning("Unsupported file format. Please upload a PDF or Word document.")

# if url_input:
#     document_text = extract_website_text(url_input)

# if document_text:
#     # Split and preprocess the text
#     sentences = split_into_sentences(document_text)
#     new_docs = preprocess_text(sentences)

#     # Vectorize the text
#     vectorizer = TfidfVectorizer(
#         analyzer="word",
#         ngram_range=(1, 2),  # Use unigrams and bigrams
#         min_df=0.01,  # Ignore very rare terms
#         max_df=0.85,  # Ignore extremely common terms
#         stop_words="english"
#     )
#     X = vectorizer.fit_transform(new_docs)

#     # DataFrame for vectorized data
#     df = pd.DataFrame(X.T.toarray(), index=vectorizer.get_feature_names_out())

#     # Search query input
#     query = st.text_input("Enter search query:")

#     # Slider to adjust the number of top results
#     top_n_results = st.slider("Number of top results to display:", min_value=1, max_value=20, value=5, step=1)

#     # Button to trigger search
#     if st.button("Search"):
#         if query:
#             # Display search results
#             results = get_similar_sentences(query, vectorizer, df, top_n=top_n_results)
#             if results:
#                 st.markdown("### Search Results")
#                 for title, sentence, score in results:
#                     st.write(f"**{title}**\n- {sentence}\n_Similarity Score: {score:.4f}_")
#             else:
#                 st.info("No similar sentences found.")
#         else:
#             st.warning("Please enter a search query.")
# else:
#     st.info("Upload a document or enter a website URL to begin.")




import re
import string
import numpy as np
import pandas as pd
import nltk
import spacy
import streamlit as st
from sklearn.feature_extraction.text import TfidfVectorizer
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
import fitz  # PyMuPDF for PDF text extraction
from docx import Document  # python-docx for Word documents
from bs4 import BeautifulSoup
import requests
import io  # To handle file-like objects for PDF processing

# --- NLTK Setup ---
nltk.download("stopwords")
nltk.download("punkt")
nltk.download("wordnet")

# --- Initialization ---
lemmatizer = WordNetLemmatizer()

# --- Text Extraction Functions ---
def extract_pdf_text(file):
    """Extract text from a PDF file."""
    doc = fitz.open(stream=file.read(), filetype="pdf")
    text = "".join([page.get_text() for page in doc])
    return text.strip()

def extract_word_text(file):
    """Extract text from a Word document."""
    doc = Document(file)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text.strip()

def extract_website_text(url):
    """Extract text from a website URL."""
    response = requests.get(url)
    soup = BeautifulSoup(response.content, "html.parser")
    paragraphs = soup.find_all("p")
    text = " ".join([para.get_text() for para in paragraphs])
    return text.strip()

# --- Text Preprocessing ---
def preprocess_text(docs):
    """Clean and lemmatize text data."""
    cleaned_docs = []
    for doc in docs:
        doc = re.sub(r'[^\x00-\x7F]+', ' ', doc)  # Remove non-ASCII characters
        doc = re.sub(r'@\w+', '', doc)  # Remove mentions
        doc = doc.lower()  # Convert to lowercase
        doc = re.sub(r'[%s]' % re.escape(string.punctuation), ' ', doc)  # Remove punctuation
        doc = re.sub(r'\d+', '', doc)  # Remove numbers
        doc = re.sub(r'\s{2,}', ' ', doc)  # Remove extra spaces
        lemmatized_doc = " ".join([lemmatizer.lemmatize(word) for word in doc.split()])
        cleaned_docs.append(lemmatized_doc)
    return cleaned_docs

# --- Utility Functions ---
def split_into_sentences(text):
    """Split text into sentences."""
    return nltk.sent_tokenize(text)

def get_similar_sentences(query, vectorizer, df, top_n=5):
    """Find similar sentences based on cosine similarity."""
    query = preprocess_text([query])  # Preprocess the query
    query_vec = vectorizer.transform(query).toarray().flatten()

    similarity_scores = {}
    for i, sentence in enumerate(new_docs):
        sentence_vec = X[i].toarray().flatten()
        numerator = np.dot(query_vec, sentence_vec)
        denominator = np.linalg.norm(query_vec) * np.linalg.norm(sentence_vec)
        similarity = numerator / denominator if denominator != 0 else 0
        similarity_scores[i] = similarity

    # Sort by similarity scores in descending order
    sorted_scores = sorted(similarity_scores.items(), key=lambda x: x[1], reverse=True)[:top_n]

    results = []
    for idx, score in sorted_scores:
        if score > 0:  # Only include sentences with positive similarity
            results.append((f"Sentence {idx + 1}", sentences[idx], score))
    return results

# --- Streamlit UI ---
# Load custom styling
st.markdown(
    """
    <style>
    .title {
        font-size: 36px;
        color: #4CAF50;
        text-align: center;
        font-weight: bold;
    }
    .subtitle {
        font-size: 18px;
        text-align: center;
        color: #555;
    }
    .menu-btn {
        background-color: #007BFF;
        color: white;
        border: none;
        border-radius: 5px;
        padding: 10px 20px;
        margin: 10px;
        font-size: 16px;
        cursor: pointer;
        text-align: center;
    }
    .menu-btn:hover {
        background-color: #0056b3;
    }
    .results {
        margin-top: 20px;
    }
    .result-item {
        background-color: #f8f9fa;
        padding: 10px;
        border-radius: 8px;
        margin-bottom: 10px;
        box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);
    }
    </style>
    """,
    unsafe_allow_html=True,
)

st.markdown('<div class="title">Document Search System</div>', unsafe_allow_html=True)
st.markdown('<div class="subtitle">Upload a document or provide a URL to perform a search query.</div>', unsafe_allow_html=True)

# --- Menu Options ---
st.markdown('<div style="text-align: center;">', unsafe_allow_html=True)
option = st.radio("Choose Input Method:", ["Upload File", "Enter URL"], horizontal=True)
st.markdown('</div>', unsafe_allow_html=True)

# --- Input Handling ---
document_text = ""
if option == "Upload File":
    uploaded_file = st.file_uploader("Upload a PDF or Word document", type=["pdf", "docx"])
    if uploaded_file:
        if uploaded_file.type == "application/pdf":
            document_text = extract_pdf_text(uploaded_file)
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            document_text = extract_word_text(uploaded_file)
        else:
            st.warning("Unsupported file format. Please upload a PDF or Word document.")
elif option == "Enter URL":
    url_input = st.text_input("Enter website URL:")
    if url_input:
        document_text = extract_website_text(url_input)

# --- Text Processing and Query ---
if document_text:
    sentences = split_into_sentences(document_text)
    new_docs = preprocess_text(sentences)

    vectorizer = TfidfVectorizer(
        analyzer="word",
        ngram_range=(1, 2),
        min_df=0.01,
        max_df=0.85,
        stop_words="english",
    )
    X = vectorizer.fit_transform(new_docs)

    query = st.text_input("Enter search query:")
    top_n_results = st.slider("Number of top results to display:", min_value=1, max_value=20, value=5, step=1)

    # Add a button to trigger the search
    if st.button("Search"):
        if query:
            results = get_similar_sentences(query, vectorizer, df=None, top_n=top_n_results)
            if results:
                st.markdown('<div class="results">', unsafe_allow_html=True)
                for title, sentence, score in results:
                    st.markdown(
                        f'<div class="result-item"><strong>{title}</strong><br><em>Similarity Score: {score:.4f}</em><br>{sentence}</div>',
                        unsafe_allow_html=True,
                    )
                st.markdown('</div>', unsafe_allow_html=True)
            else:
                st.info("No similar sentences found.")
        else:
            st.warning("Please enter a search query.")
else:
    st.info("Upload a document or enter a website URL to begin.")
