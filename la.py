



# import re
# import string
# import numpy as np
# import pandas as pd
# import nltk
# import spacy
# import streamlit as st
# from sklearn.feature_extraction.text import TfidfVectorizer
# from nltk.corpus import stopwords
# from nltk.stem import WordNetLemmatizer
# import fitz  # PyMuPDF for PDF text extraction
# from docx import Document  # python-docx for Word documents
# import requests
# from bs4 import BeautifulSoup
# import requests
# # import io
# import io  # To handle file-like object for PDF


# # NLTK setup
# nltk.download("stopwords")
# nltk.download("punkt")
# nltk.download("wordnet")

# # Lemmatization and Preprocessing
# lemmer = WordNetLemmatizer()

# # Function to extract text from PDF (handling file-like object)
# def extract_pdf_text(file):
#     doc = fitz.open(stream=file.read(), filetype="pdf")  # Directly use the file buffer with PyMuPDF
#     text = ""
#     for page in doc:
#         text += page.get_text()
#     return text

# # Function to extract text from Word document
# def extract_word_text(file):
#     doc = Document(file)
#     text = ""
#     for para in doc.paragraphs:
#         text += para.text + "\n"
#     return text

# # Function to extract text from URL (Website)
# def extract_website_text(url):
#     response = requests.get(url)
#     soup = BeautifulSoup(response.content, "html.parser")
#     paragraphs = soup.find_all("p")
#     text = ""
#     for para in paragraphs:
#         text += para.get_text() + "\n"
#     return text

# # Preprocessing function to clean and lemmatize text
# def preprocess_text(docs):
#     cleaned_docs = []
#     for doc in docs:
#         doc = re.sub(r'[^\x00-\x7F]+', ' ', doc)  # Remove non-ASCII characters
#         doc = re.sub(r'@\w+', '', doc)  # Remove Mentions
#         doc = doc.lower()  # Convert to lowercase
#         doc = re.sub(r'[%s]' % re.escape(string.punctuation), ' ', doc)  # Clean punctuation
#         doc = re.sub(r'[0-9]', '', doc)  # Remove numbers
#         doc = re.sub(r'\s{2,}', ' ', doc)  # Remove extra spaces
#         cleaned_docs.append(' '.join([lemmer.lemmatize(word) for word in doc.split()]))
#     return cleaned_docs

# # Split text into sentences
# def split_into_sentences(text):
#     return nltk.sent_tokenize(text)

# # Streamlit UI
# st.title("Document Search System")

# # File Upload for PDF, Word, or Text
# uploaded_file = st.file_uploader("Upload a PDF or Word document", type=["pdf", "docx"])

# # URL Input for Website
# url_input = st.text_input("Enter website URL (for scraping)")

# if uploaded_file:
#     if uploaded_file.type == "application/pdf":
#         document_text = extract_pdf_text(uploaded_file)
#     elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#         document_text = extract_word_text(uploaded_file)
#     else:
#         st.warning("Unsupported file format")
# else:
#     document_text = ""
    
# if url_input:
#     document_text = extract_website_text(url_input)

# if document_text:
#     # Split text into sentences and preprocess
#     sentences = split_into_sentences(document_text)
#     new_docs = preprocess_text(sentences)

#     # Vectorizing the document with adjusted min_df and max_df
#     vectorizer = TfidfVectorizer(
#         analyzer='word', 
#         ngram_range=(1, 2),  # Using unigrams and bigrams
#         min_df=0.0,  # Lowered min_df to allow more terms
#         max_df=1.0,  # Raised max_df to avoid excluding too many common words
#         max_features=10000, 
#         lowercase=True, 
#         stop_words='english'
#     )
#     X = vectorizer.fit_transform(new_docs)
    
#     # Ensure that terms are present after vectorization
#     if X.shape[1] == 0:
#         st.write("No terms were extracted from the document. Try using a different document or adjusting the min_df/max_df parameters.")
#     else:
#         df = pd.DataFrame(X.T.toarray())
        
#         # Function to search for similar sentences
#         def get_similar_sentences(query, df):
#             query = [query]
#             query_vec = vectorizer.transform(query).toarray().reshape(df.shape[0],)
            
#             sim = {}
#             for i in range(len(new_docs)):
#                 sim[i] = np.dot(df.loc[:, i].values, query_vec) / (np.linalg.norm(df.loc[:, i]) * np.linalg.norm(query_vec))
            
#             sim_sorted = sorted(sim.items(), key=lambda x: x[1], reverse=True)[:5]
            
#             result = []
#             for i, v in sim_sorted:
#                 if v != 0.0:
#                     result.append((f"Sentence {i + 1}", new_docs[i], v))  # For simplicity, just use sentence index
#             return result

#         # Display the input search query
#         query = st.text_input("Enter search query:")
#         if query:
#             results = get_similar_sentences(query, df)
#             if results:
#                 st.subheader("Top Results:")
#                 for title, doc, score in results:
#                     st.write(f"**Title:** {title}")
#                     st.write(f"**Similarity Score:** {score}")
#                     st.write(f"**Sentence:** {doc}")
#                     st.write("-" * 100)
#             else:
#                 st.write("No results found.")
# else:
#     st.write("Upload a document or enter a website URL to begin.")






# import re
# import string
# import numpy as np
# import pandas as pd
# import nltk
# import spacy
# import streamlit as st
# from sklearn.feature_extraction.text import TfidfVectorizer
# from nltk.corpus import stopwords
# from nltk.stem import WordNetLemmatizer
# import fitz  # PyMuPDF for PDF text extraction
# from docx import Document  # python-docx for Word documents
# from bs4 import BeautifulSoup
# import requests
# import io  # To handle file-like objects for PDF processing

# # NLTK setup
# nltk.download("stopwords")
# nltk.download("punkt")
# nltk.download("wordnet")

# # Lemmatization and Preprocessing
# lemmer = WordNetLemmatizer()

# # Function to extract text from PDF
# def extract_pdf_text(file):
#     doc = fitz.open(stream=file.read(), filetype="pdf")
#     text = ""
#     for page in doc:
#         text += page.get_text()
#     return text.strip()

# # Function to extract text from Word document
# def extract_word_text(file):
#     doc = Document(file)
#     text = ""
#     for para in doc.paragraphs:
#         text += para.text + "\n"
#     return text.strip()

# # Function to extract text from a website URL
# def extract_website_text(url):
#     response = requests.get(url)
#     soup = BeautifulSoup(response.content, "html.parser")
#     paragraphs = soup.find_all("p")
#     text = " ".join([para.get_text() for para in paragraphs])
#     return text.strip()

# # Function to preprocess and lemmatize text
# def preprocess_text(docs):
#     cleaned_docs = []
#     for doc in docs:
#         doc = re.sub(r'[^\x00-\x7F]+', ' ', doc)  # Remove non-ASCII characters
#         doc = re.sub(r'@\w+', '', doc)  # Remove mentions
#         doc = doc.lower()  # Convert to lowercase
#         doc = re.sub(r'[%s]' % re.escape(string.punctuation), ' ', doc)  # Remove punctuation
#         doc = re.sub(r'\d+', '', doc)  # Remove numbers
#         doc = re.sub(r'\s{2,}', ' ', doc)  # Remove extra spaces
#         lemmatized_doc = " ".join([lemmer.lemmatize(word) for word in doc.split()])
#         cleaned_docs.append(lemmatized_doc)
#     return cleaned_docs

# # Function to split text into sentences
# def split_into_sentences(text):
#     return nltk.sent_tokenize(text)

# # Streamlit UI
# st.title("Document Search System")

# # File upload for PDF or Word documents
# uploaded_file = st.file_uploader("Upload a PDF or Word document", type=["pdf", "docx"])

# # URL input for scraping text
# url_input = st.text_input("Enter website URL (for scraping)")

# document_text = ""

# # Extract text based on the input type
# if uploaded_file:
#     if uploaded_file.type == "application/pdf":
#         document_text = extract_pdf_text(uploaded_file)
#     elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#         document_text = extract_word_text(uploaded_file)
#     else:
#         st.warning("Unsupported file format. Please upload a PDF or Word document.")

# if url_input:
#     document_text = extract_website_text(url_input)

# if document_text:
#     # Split text into sentences and preprocess
#     sentences = split_into_sentences(document_text)
#     new_docs = preprocess_text(sentences)

#     # Vectorize the document
#     vectorizer = TfidfVectorizer(
#         analyzer="word",
#         ngram_range=(1, 2),  # Use unigrams and bigrams
#         min_df=0.01,  # Ignore very rare terms
#         max_df=0.85,  # Ignore extremely common terms
#         stop_words="english"
#     )
#     X = vectorizer.fit_transform(new_docs)

#     # Create DataFrame for vectorized data
#     df = pd.DataFrame(X.T.toarray(), index=vectorizer.get_feature_names_out())

#     # Function to search for similar sentences
#     def get_similar_sentences(query, vectorizer, df):
#         query = preprocess_text([query])  # Preprocess the query
#         query_vec = vectorizer.transform(query).toarray().flatten()

#         similarity_scores = {}
#         for i, sentence in enumerate(new_docs):
#             sentence_vec = X[i].toarray().flatten()
#             # Calculate cosine similarity
#             numerator = np.dot(query_vec, sentence_vec)
#             denominator = np.linalg.norm(query_vec) * np.linalg.norm(sentence_vec)
#             similarity = numerator / denominator if denominator != 0 else 0
#             similarity_scores[i] = similarity

#         # Sort by similarity scores in descending order
#         sorted_scores = sorted(similarity_scores.items(), key=lambda x: x[1], reverse=True)[:5]

#         results = []
#         for idx, score in sorted_scores:
#             if score > 0:  # Only include sentences with positive similarity
#                 results.append((f"Sentence {idx + 1}", sentences[idx], score))
#         return results

#     # Search query input
#     query = st.text_input("Enter search query:")
#     if query:
#         results = get_similar_sentences(query, vectorizer, df)
#         if results:
#             st.subheader("Top Results:")
#             for title, sentence, score in results:
#                 st.write(f"**{title}**")
#                 st.write(f"**Similarity Score:** {score:.4f}")
#                 st.write(f"**Sentence:** {sentence}")
#                 st.write("-" * 100)
#         else:
#             st.write("No similar sentences found.")
# else:
#     st.write("Upload a document or enter a website URL to begin.")





# import re
# import string
# import numpy as np
# import pandas as pd
# import nltk
# import spacy
# import streamlit as st
# from sklearn.feature_extraction.text import TfidfVectorizer
# from nltk.corpus import stopwords
# from nltk.stem import WordNetLemmatizer
# import fitz  # PyMuPDF for PDF text extraction
# from docx import Document  # python-docx for Word documents
# from bs4 import BeautifulSoup
# import requests
# import io  # To handle file-like objects for PDF processing

# # NLTK setup
# nltk.download("stopwords")
# nltk.download("punkt")
# nltk.download("wordnet")

# # Lemmatization and Preprocessing
# lemmer = WordNetLemmatizer()

# # Function to extract text from PDF
# def extract_pdf_text(file):
#     doc = fitz.open(stream=file.read(), filetype="pdf")
#     text = ""
#     for page in doc:
#         text += page.get_text()
#     return text.strip()

# # Function to extract text from Word document
# def extract_word_text(file):
#     doc = Document(file)
#     text = ""
#     for para in doc.paragraphs:
#         text += para.text + "\n"
#     return text.strip()

# # Function to extract text from a website URL
# def extract_website_text(url):
#     response = requests.get(url)
#     soup = BeautifulSoup(response.content, "html.parser")
#     paragraphs = soup.find_all("p")
#     text = " ".join([para.get_text() for para in paragraphs])
#     return text.strip()

# # Function to preprocess and lemmatize text
# def preprocess_text(docs):
#     cleaned_docs = []
#     for doc in docs:
#         doc = re.sub(r'[^\x00-\x7F]+', ' ', doc)  # Remove non-ASCII characters
#         doc = re.sub(r'@\w+', '', doc)  # Remove mentions
#         doc = doc.lower()  # Convert to lowercase
#         doc = re.sub(r'[%s]' % re.escape(string.punctuation), ' ', doc)  # Remove punctuation
#         doc = re.sub(r'\d+', '', doc)  # Remove numbers
#         doc = re.sub(r'\s{2,}', ' ', doc)  # Remove extra spaces
#         lemmatized_doc = " ".join([lemmer.lemmatize(word) for word in doc.split()])
#         cleaned_docs.append(lemmatized_doc)
#     return cleaned_docs

# # Function to split text into sentences
# def split_into_sentences(text):
#     return nltk.sent_tokenize(text)

# # Streamlit UI
# st.title("Document Search System")

# # File upload for PDF or Word documents
# uploaded_file = st.file_uploader("Upload a PDF or Word document", type=["pdf", "docx"])

# # URL input for scraping text
# url_input = st.text_input("Enter website URL (for scraping)")

# document_text = ""

# # Extract text based on the input type
# if uploaded_file:
#     if uploaded_file.type == "application/pdf":
#         document_text = extract_pdf_text(uploaded_file)
#     elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#         document_text = extract_word_text(uploaded_file)
#     else:
#         st.warning("Unsupported file format. Please upload a PDF or Word document.")

# if url_input:
#     document_text = extract_website_text(url_input)

# if document_text:
#     # Split text into sentences and preprocess
#     sentences = split_into_sentences(document_text)
#     new_docs = preprocess_text(sentences)

#     # Vectorize the document
#     vectorizer = TfidfVectorizer(
#         analyzer="word",
#         ngram_range=(1, 2),  # Use unigrams and bigrams
#         min_df=0.01,  # Ignore very rare terms
#         max_df=0.85,  # Ignore extremely common terms
#         stop_words="english"
#     )
#     X = vectorizer.fit_transform(new_docs)

#     # Create DataFrame for vectorized data
#     df = pd.DataFrame(X.T.toarray(), index=vectorizer.get_feature_names_out())

#     # Function to search for similar sentences
#     def get_similar_sentences(query, vectorizer, df, top_n=5):
#         query = preprocess_text([query])  # Preprocess the query
#         query_vec = vectorizer.transform(query).toarray().flatten()

#         similarity_scores = {}
#         for i, sentence in enumerate(new_docs):
#             sentence_vec = X[i].toarray().flatten()
#             # Calculate cosine similarity
#             numerator = np.dot(query_vec, sentence_vec)
#             denominator = np.linalg.norm(query_vec) * np.linalg.norm(sentence_vec)
#             similarity = numerator / denominator if denominator != 0 else 0
#             similarity_scores[i] = similarity

#         # Sort by similarity scores in descending order
#         sorted_scores = sorted(similarity_scores.items(), key=lambda x: x[1], reverse=True)[:top_n]

#         results = []
#         for idx, score in sorted_scores:
#             if score > 0:  # Only include sentences with positive similarity
#                 results.append((f"Sentence {idx + 1}", sentences[idx], score))
#         return results

#     # Search query input
#     query = st.text_input("Enter search query:")

#     # Slider to adjust the number of top results
#     top_n_results = st.slider("Number of top results to display:", min_value=1, max_value=10, value=5, step=1)

#     if query:
#         results = get_similar_sentences(query, vectorizer, df, top_n=top_n_results)
#         if results:
#             st.subheader("Top Results:")
#             for title, sentence, score in results:
#                 st.write(f"**{title}**")
#                 st.write(f"**Similarity Score:** {score:.4f}")
#                 st.write(f"**Sentence:** {sentence}")
#                 st.write("-" * 100)
#         else:
#             st.write("No similar sentences found.")
# else:
#     st.write("Upload a document or enter a website URL to begin.")





# import re
# import string
# import numpy as np
# import pandas as pd
# import nltk
# import spacy
# import streamlit as st
# from sklearn.feature_extraction.text import TfidfVectorizer
# from nltk.corpus import stopwords
# from nltk.stem import WordNetLemmatizer
# import fitz  # PyMuPDF for PDF text extraction
# from docx import Document  # python-docx for Word documents
# from bs4 import BeautifulSoup
# import requests
# import io  # To handle file-like objects for PDF processing

# # --- NLTK Setup ---
# nltk.download("stopwords")
# nltk.download("punkt")
# nltk.download("wordnet")

# # --- Initialization ---
# lemmatizer = WordNetLemmatizer()

# # --- Text Extraction Functions ---
# def extract_pdf_text(file):
#     """Extract text from a PDF file."""
#     doc = fitz.open(stream=file.read(), filetype="pdf")
#     text = "".join([page.get_text() for page in doc])
#     return text.strip()

# def extract_word_text(file):
#     """Extract text from a Word document."""
#     doc = Document(file)
#     text = "\n".join([para.text for para in doc.paragraphs])
#     return text.strip()

# def extract_website_text(url):
#     """Extract text from a website URL."""
#     response = requests.get(url)
#     soup = BeautifulSoup(response.content, "html.parser")
#     paragraphs = soup.find_all("p")
#     text = " ".join([para.get_text() for para in paragraphs])
#     return text.strip()

# # --- Text Preprocessing ---
# def preprocess_text(docs):
#     """Clean and lemmatize text data."""
#     cleaned_docs = []
#     for doc in docs:
#         doc = re.sub(r'[^\x00-\x7F]+', ' ', doc)  # Remove non-ASCII characters
#         doc = re.sub(r'@\w+', '', doc)  # Remove mentions
#         doc = doc.lower()  # Convert to lowercase
#         doc = re.sub(r'[%s]' % re.escape(string.punctuation), ' ', doc)  # Remove punctuation
#         doc = re.sub(r'\d+', '', doc)  # Remove numbers
#         doc = re.sub(r'\s{2,}', ' ', doc)  # Remove extra spaces
#         lemmatized_doc = " ".join([lemmatizer.lemmatize(word) for word in doc.split()])
#         cleaned_docs.append(lemmatized_doc)
#     return cleaned_docs

# # --- Utility Functions ---
# def split_into_sentences(text):
#     """Split text into sentences."""
#     return nltk.sent_tokenize(text)

# def get_similar_sentences(query, vectorizer, df, top_n=5):
#     """Find similar sentences based on cosine similarity."""
#     query = preprocess_text([query])  # Preprocess the query
#     query_vec = vectorizer.transform(query).toarray().flatten()

#     similarity_scores = {}
#     for i, sentence in enumerate(new_docs):
#         sentence_vec = X[i].toarray().flatten()
#         numerator = np.dot(query_vec, sentence_vec)
#         denominator = np.linalg.norm(query_vec) * np.linalg.norm(sentence_vec)
#         similarity = numerator / denominator if denominator != 0 else 0
#         similarity_scores[i] = similarity

#     # Sort by similarity scores in descending order
#     sorted_scores = sorted(similarity_scores.items(), key=lambda x: x[1], reverse=True)[:top_n]

#     results = []
#     for idx, score in sorted_scores:
#         if score > 0:  # Only include sentences with positive similarity
#             results.append((f"Sentence {idx + 1}", sentences[idx], score))
#     return results

# # --- Streamlit UI ---
# st.title("Document Search System")
# st.markdown("Upload a document or provide a URL to perform a search query.")

# # File upload for PDF or Word documents
# uploaded_file = st.file_uploader("Upload a PDF or Word document", type=["pdf", "docx"])

# # URL input for scraping text
# url_input = st.text_input("Enter website URL (for scraping)")

# # Initialize document text
# document_text = ""

# # Extract text from the uploaded file or URL
# if uploaded_file:
#     if uploaded_file.type == "application/pdf":
#         document_text = extract_pdf_text(uploaded_file)
#     elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#         document_text = extract_word_text(uploaded_file)
#     else:
#         st.warning("Unsupported file format. Please upload a PDF or Word document.")

# if url_input:
#     document_text = extract_website_text(url_input)

# if document_text:
#     # Split and preprocess the text
#     sentences = split_into_sentences(document_text)
#     new_docs = preprocess_text(sentences)

#     # Vectorize the text
#     vectorizer = TfidfVectorizer(
#         analyzer="word",
#         ngram_range=(1, 2),  # Use unigrams and bigrams
#         min_df=0.01,  # Ignore very rare terms
#         max_df=0.85,  # Ignore extremely common terms
#         stop_words="english"
#     )
#     X = vectorizer.fit_transform(new_docs)

#     # DataFrame for vectorized data
#     df = pd.DataFrame(X.T.toarray(), index=vectorizer.get_feature_names_out())

#     # Search query input
#     query = st.text_input("Enter search query:")

#     # Slider to adjust the number of top results
#     top_n_results = st.slider("Number of top results to display:", min_value=1, max_value=20, value=5, step=1)

#     if query:
#         # Display search results
#         results = get_similar_sentences(query, vectorizer, df, top_n=top_n_results)
#         if results:
#             st.subheader("Top Results:")
#             for title, sentence, score in results:
#                 st.write(f"**{title}**")
#                 st.write(f"**Similarity Score:** {score:.4f}")
#                 st.write(f"**Sentence:** {sentence}")
#                 st.write("-" * 100)
#         else:
#             st.info("No similar sentences found.")
# else:
#     st.info("Upload a document or enter a website URL to begin.")














# import re
# import string
# import numpy as np
# import pandas as pd
# import nltk
# import spacy
# import streamlit as st
# from sklearn.feature_extraction.text import TfidfVectorizer
# from nltk.corpus import stopwords
# from nltk.stem import WordNetLemmatizer
# import fitz  # PyMuPDF for PDF text extraction
# from docx import Document  # python-docx for Word documents
# from bs4 import BeautifulSoup
# import requests
# import io  # To handle file-like objects for PDF processing

# # --- NLTK Setup ---
# nltk.download("stopwords")
# nltk.download("punkt")
# nltk.download("wordnet")

# # --- Initialization ---
# lemmatizer = WordNetLemmatizer()

# # --- Text Extraction Functions ---
# def extract_pdf_text(file):
#     """Extract text from a PDF file."""
#     doc = fitz.open(stream=file.read(), filetype="pdf")
#     text = "".join([page.get_text() for page in doc])
#     return text.strip()

# def extract_word_text(file):
#     """Extract text from a Word document."""
#     doc = Document(file)
#     text = "\n".join([para.text for para in doc.paragraphs])
#     return text.strip()

# def extract_website_text(url):
#     """Extract text from a website URL."""
#     response = requests.get(url)
#     soup = BeautifulSoup(response.content, "html.parser")
#     paragraphs = soup.find_all("p")
#     text = " ".join([para.get_text() for para in paragraphs])
#     return text.strip()

# # --- Text Preprocessing ---
# def preprocess_text(docs):
#     """Clean and lemmatize text data."""
#     cleaned_docs = []
#     for doc in docs:
#         doc = re.sub(r'[^\x00-\x7F]+', ' ', doc)  # Remove non-ASCII characters
#         doc = re.sub(r'@\w+', '', doc)  # Remove mentions
#         doc = doc.lower()  # Convert to lowercase
#         doc = re.sub(r'[%s]' % re.escape(string.punctuation), ' ', doc)  # Remove punctuation
#         doc = re.sub(r'\d+', '', doc)  # Remove numbers
#         doc = re.sub(r'\s{2,}', ' ', doc)  # Remove extra spaces
#         lemmatized_doc = " ".join([lemmatizer.lemmatize(word) for word in doc.split()])
#         cleaned_docs.append(lemmatized_doc)
#     return cleaned_docs

# # --- Utility Functions ---
# def split_into_sentences(text):
#     """Split text into sentences."""
#     return nltk.sent_tokenize(text)

# def get_similar_sentences(query, vectorizer, df, top_n=5):
#     """Find similar sentences based on cosine similarity."""
#     query = preprocess_text([query])  # Preprocess the query
#     query_vec = vectorizer.transform(query).toarray().flatten()

#     similarity_scores = {}
#     for i, sentence in enumerate(new_docs):
#         sentence_vec = X[i].toarray().flatten()
#         numerator = np.dot(query_vec, sentence_vec)
#         denominator = np.linalg.norm(query_vec) * np.linalg.norm(sentence_vec)
#         similarity = numerator / denominator if denominator != 0 else 0
#         similarity_scores[i] = similarity

#     # Sort by similarity scores in descending order
#     sorted_scores = sorted(similarity_scores.items(), key=lambda x: x[1], reverse=True)[:top_n]

#     results = []
#     for idx, score in sorted_scores:
#         if score > 0:  # Only include sentences with positive similarity
#             results.append((f"Sentence {idx + 1}", sentences[idx], score))
#     return results

# # --- Streamlit UI --- 
# st.markdown("""
#     <style>
#     .title {
#         text-align: center;
#         color: #2d2d2d;
#         font-size: 30px;
#         font-weight: bold;
#         padding: 20px;
#     }
#     .subtitle {
#         text-align: center;
#         font-size: 18px;
#         color: #5a5a5a;
#     }
#     .file-uploader, .url-input {
#         display: flex;
#         justify-content: center;
#         margin: 20px 0;
#     }
#     .file-uploader button, .url-input input {
#         font-size: 16px;
#         padding: 10px;
#         border-radius: 5px;
#         border: 2px solid #ddd;
#         background-color: #f0f0f0;
#     }
#     .file-uploader button:hover, .url-input input:hover {
#         background-color: #e0e0e0;
#         cursor: pointer;
#     }
#     .results {
#         background-color: #f9f9f9;
#         padding: 20px;
#         border-radius: 8px;
#         margin-top: 20px;
#         border: 1px solid #ddd;
#     }
#     .result-item {
#         margin-bottom: 15px;
#     }
#     </style>
# """, unsafe_allow_html=True)

# st.markdown('<div class="title">Document Search System</div>', unsafe_allow_html=True)
# st.markdown('<div class="subtitle">Upload a document or provide a URL to perform a search query.</div>', unsafe_allow_html=True)

# # File upload for PDF or Word documents
# uploaded_file = st.file_uploader("Upload a PDF or Word document", type=["pdf", "docx"], key="file_uploader")

# # URL input for scraping text
# url_input = st.text_input("Enter website URL (for scraping)", key="url_input")

# # Initialize document text
# document_text = ""

# # Extract text from the uploaded file or URL
# if uploaded_file:
#     if uploaded_file.type == "application/pdf":
#         document_text = extract_pdf_text(uploaded_file)
#     elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#         document_text = extract_word_text(uploaded_file)
#     else:
#         st.warning("Unsupported file format. Please upload a PDF or Word document.")

# if url_input:
#     document_text = extract_website_text(url_input)

# if document_text:
#     # Split and preprocess the text
#     sentences = split_into_sentences(document_text)
#     new_docs = preprocess_text(sentences)

#     # Vectorize the text
#     vectorizer = TfidfVectorizer(
#         analyzer="word",
#         ngram_range=(1, 2),  # Use unigrams and bigrams
#         min_df=0.01,  # Ignore very rare terms
#         max_df=0.85,  # Ignore extremely common terms
#         stop_words="english"
#     )
#     X = vectorizer.fit_transform(new_docs)

#     # DataFrame for vectorized data
#     df = pd.DataFrame(X.T.toarray(), index=vectorizer.get_feature_names_out())

#     # Search query input
#     query = st.text_input("Enter search query:")

#     # Slider to adjust the number of top results
#     top_n_results = st.slider("Number of top results to display:", min_value=1, max_value=20, value=5, step=1)

#     if query:
#         # Display search results
#         results = get_similar_sentences(query, vectorizer, df, top_n=top_n_results)
#         if results:
#             st.markdown('<div class="results">', unsafe_allow_html=True)
#             for title, sentence, score in results:
#                 st.markdown(f'<div class="result-item"><strong>{title}</strong><br><em>Similarity Score: {score:.4f}</em><br>{sentence}</div>', unsafe_allow_html=True)
#             st.markdown('</div>', unsafe_allow_html=True)
#         else:
#             st.info("No similar sentences found.")
# else:
#     st.info("Upload a document or enter a website URL to begin.")






import re
import string
import numpy as np
import pandas as pd
import nltk
import spacy
import streamlit as st
from sklearn.feature_extraction.text import TfidfVectorizer
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
import fitz  # PyMuPDF for PDF text extraction
from docx import Document  # python-docx for Word documents
from bs4 import BeautifulSoup
import requests
import io  # To handle file-like objects for PDF processing

# --- NLTK Setup ---
nltk.download("stopwords")
nltk.download("punkt")
nltk.download("wordnet")

# --- Initialization ---
lemmatizer = WordNetLemmatizer()

# --- Text Extraction Functions ---
def extract_pdf_text(file):
    """Extract text from a PDF file."""
    doc = fitz.open(stream=file.read(), filetype="pdf")
    text = "".join([page.get_text() for page in doc])
    return text.strip()

def extract_word_text(file):
    """Extract text from a Word document."""
    doc = Document(file)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text.strip()

def extract_website_text(url):
    """Extract text from a website URL."""
    response = requests.get(url)
    soup = BeautifulSoup(response.content, "html.parser")
    paragraphs = soup.find_all("p")
    text = " ".join([para.get_text() for para in paragraphs])
    return text.strip()

# --- Text Preprocessing ---
def preprocess_text(docs):
    """Clean and lemmatize text data."""
    cleaned_docs = []
    for doc in docs:
        doc = re.sub(r'[^\x00-\x7F]+', ' ', doc)  # Remove non-ASCII characters
        doc = re.sub(r'@\w+', '', doc)  # Remove mentions
        doc = doc.lower()  # Convert to lowercase
        doc = re.sub(r'[%s]' % re.escape(string.punctuation), ' ', doc)  # Remove punctuation
        doc = re.sub(r'\d+', '', doc)  # Remove numbers
        doc = re.sub(r'\s{2,}', ' ', doc)  # Remove extra spaces
        lemmatized_doc = " ".join([lemmatizer.lemmatize(word) for word in doc.split()])
        cleaned_docs.append(lemmatized_doc)
    return cleaned_docs

# --- Utility Functions ---
def split_into_sentences(text):
    """Split text into sentences."""
    return nltk.sent_tokenize(text)

def get_similar_sentences(query, vectorizer, df, top_n=5):
    """Find similar sentences based on cosine similarity."""
    query = preprocess_text([query])  # Preprocess the query
    query_vec = vectorizer.transform(query).toarray().flatten()

    similarity_scores = {}
    for i, sentence in enumerate(new_docs):
        sentence_vec = X[i].toarray().flatten()
        numerator = np.dot(query_vec, sentence_vec)
        denominator = np.linalg.norm(query_vec) * np.linalg.norm(sentence_vec)
        similarity = numerator / denominator if denominator != 0 else 0
        similarity_scores[i] = similarity

    # Sort by similarity scores in descending order
    sorted_scores = sorted(similarity_scores.items(), key=lambda x: x[1], reverse=True)[:top_n]

    results = []
    for idx, score in sorted_scores:
        if score > 0:  # Only include sentences with positive similarity
            results.append((f"Sentence {idx + 1}", sentences[idx], score))
    return results

# --- Streamlit UI --- 
# Load the CSS file for custom styling
with open("style.css", "r") as file:
    st.markdown(f"<style>{file.read()}</style>", unsafe_allow_html=True)

st.markdown('<div class="title">Document Search System</div>', unsafe_allow_html=True)
st.markdown('<div class="subtitle">Upload a document or provide a URL to perform a search query.</div>', unsafe_allow_html=True)

# File upload for PDF or Word documents
uploaded_file = st.file_uploader("Upload a PDF or Word document", type=["pdf", "docx"], key="file_uploader")

# URL input for scraping text
url_input = st.text_input("Enter website URL (for scraping)", key="url_input")

# Initialize document text
document_text = ""

# Extract text from the uploaded file or URL
if uploaded_file:
    if uploaded_file.type == "application/pdf":
        document_text = extract_pdf_text(uploaded_file)
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        document_text = extract_word_text(uploaded_file)
    else:
        st.warning("Unsupported file format. Please upload a PDF or Word document.")

if url_input:
    document_text = extract_website_text(url_input)

if document_text:
    # Split and preprocess the text
    sentences = split_into_sentences(document_text)
    new_docs = preprocess_text(sentences)

    # Vectorize the text
    vectorizer = TfidfVectorizer(
        analyzer="word",
        ngram_range=(1, 2),  # Use unigrams and bigrams
        min_df=0.01,  # Ignore very rare terms
        max_df=0.85,  # Ignore extremely common terms
        stop_words="english"
    )
    X = vectorizer.fit_transform(new_docs)

    # DataFrame for vectorized data
    df = pd.DataFrame(X.T.toarray(), index=vectorizer.get_feature_names_out())

    # Search query input
    query = st.text_input("Enter search query:")

    # Slider to adjust the number of top results
    top_n_results = st.slider("Number of top results to display:", min_value=1, max_value=20, value=5, step=1)

    if query:
        # Display search results
        results = get_similar_sentences(query, vectorizer, df, top_n=top_n_results)
        if results:
            st.markdown('<div class="results">', unsafe_allow_html=True)
            for title, sentence, score in results:
                st.markdown(f'<div class="result-item"><strong>{title}</strong><br><em>Similarity Score: {score:.4f}</em><br>{sentence}</div>', unsafe_allow_html=True)
            st.markdown('</div>', unsafe_allow_html=True)
        else:
            st.info("No similar sentences found.")
else:
    st.info("Upload a document or enter a website URL to begin.")
